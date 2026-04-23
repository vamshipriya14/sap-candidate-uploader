"""
Resume Upload Web App — Backend API
Deploys to Render.com (free tier) — no VPS/SSH needed.
Uses the real Python resume parser (spaCy + pdfplumber + phonenumbers).
"""

import os
import re
import io
import hashlib
import traceback
from datetime import datetime, timezone
from functools import wraps

from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
import requests

# Resume parsing
import pdfplumber
import docx as python_docx
import spacy
import phonenumbers

load_dotenv()

app = Flask(__name__)

# Allow requests from any origin (GoDaddy subdomain will call this)
CORS(app, resources={r"/api/*": {"origins": "*"}})

# ─── Config ─────────────────────────────────────────────────────────────────
SUPABASE_URL   = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY   = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")
GH_REPO        = os.getenv("GH_REPO", "")
GH_TOKEN       = os.getenv("GH_TOKEN", "")
GH_EVENT       = os.getenv("GH_EVENT_TYPE", "resume-form-submitted")
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "*")  # e.g. "https://upload.yourdomain.com"

try:
    nlp = spacy.load("en_core_web_sm")
    NLP_LOADED = True
except Exception:
    NLP_LOADED = False

def _headers(is_binary=False):
    h = {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Prefer": "return=representation",
    }
    if not is_binary:
        h["Content-Type"] = "application/json"
    else:
        h["Content-Type"] = "application/octet-stream"
    return h


# ─── Resume Parser (same logic as src/resume_parser.py) ─────────────────────

INVALID_WORDS = {
    "resume", "cv", "professional", "summary", "profile", "skills",
    "experience", "business", "operations", "management", "core",
    "competencies", "personal", "details", "id", "objective",
    "education", "contact", "address", "declaration", "references",
    "linkedin", "github", "portfolio"
}

COUNTRY_MAP = {
    "+91":  ("+91",  "India"),
    "+1":   ("+1",   "United States"),
    "+44":  ("+44",  "United Kingdom"),
    "+61":  ("+61",  "Australia"),
    "+971": ("+971", "United Arab Emirates"),
    "+65":  ("+65",  "Singapore"),
    "+60":  ("+60",  "Malaysia"),
    "+49":  ("+49",  "Germany"),
    "+33":  ("+33",  "France"),
    "+81":  ("+81",  "Japan"),
}


def _extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    text = ""
    if filename.lower().endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    elif filename.lower().endswith(".docx"):
        doc = python_docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    return text


def _extract_email(text: str) -> str:
    text = re.sub(r'(\S+)\s*@\s*(\S+)', r'\1@\2', text)
    emails = re.findall(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', text)
    return emails[0] if emails else ""


def _extract_phone(text: str) -> str:
    for region in ("IN", None):
        for match in phonenumbers.PhoneNumberMatcher(text, region):
            return phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.E164)
    return ""


def _extract_country(phone: str):
    for code, value in COUNTRY_MAP.items():
        if phone.startswith(code):
            return value
    return "+91", "India"


def _extract_name(text: str, email: str = None):
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    def clean_line(line):
        line = re.sub(r'\S+@\S+', '', line)
        line = re.sub(r'\+?\d[\d\s\-]{8,}', '', line)
        line = re.sub(r'(?i)\b(email|mail|contact|phone|mobile)\b[:\-]*', '', line)
        line = re.sub(r'[^\w\s]', ' ', line)
        return re.sub(r'\s+', ' ', line).strip()

    def is_valid(words):
        if not (2 <= len(words) <= 4):
            return False
        return all(
            not any(c.isdigit() for c in w) and w.lower() not in INVALID_WORDS
            for w in words
        )

    def split_name(words):
        words = [w.capitalize() for w in words]
        if len(words[0]) == 1:
            return " ".join(words[1:]), words[0]
        if len(words) == 2:
            return words[0], words[1]
        return " ".join(words[:-1]), words[-1]

    spacy_names = set()
    if NLP_LOADED:
        spacy_doc = nlp(text[:1000])
        spacy_names = {ent.text.lower() for ent in spacy_doc.ents if ent.label_ == "PERSON"}

    def spacy_match(name):
        name_lower = name.lower()
        return any(name_lower in s or s in name_lower for s in spacy_names)

    candidates = []
    for idx, line in enumerate(lines[:20]):
        words = clean_line(line).split()
        if is_valid(words):
            candidates.append((words, idx))

    for i, line in enumerate(lines):
        if "@" in line:
            for target_line in ([lines[i - 1]] if i > 0 else []) + [line]:
                words = clean_line(target_line).split()
                if is_valid(words):
                    candidates.append((words, i))

    email_username = re.sub(r'[^a-z]', '', email.split("@")[0].lower()) if email else ""
    best_words, best_score = None, -1

    for words, position in candidates:
        name = " ".join(words)
        score = 2
        if position < 5:
            score += 2
        elif position < 10:
            score += 1
        if 2 <= len(words) <= 3:
            score += 2
        if email_username and any(w.lower() in email_username for w in words):
            score += 3
        if spacy_match(name):
            score += 3
        if score > best_score:
            best_score, best_words = score, words

    if best_words:
        first, last = split_name(best_words)
        confidence = "high" if best_score >= 7 else "medium" if best_score >= 4 else "low"
        return first, last, confidence

    return "", "", "low"


def parse_resume_bytes(file_bytes: bytes, filename: str) -> dict:
    text = _extract_text_from_bytes(file_bytes, filename)
    email = _extract_email(text)
    first, last, confidence = _extract_name(text, email)
    phone = _extract_phone(text)
    code, country = _extract_country(phone)
    return {
        "first_name":   first,
        "last_name":    last,
        "email":        email,
        "phone":        phone,
        "country_code": code,
        "country":      country,
        "confidence":   confidence,
    }


# ─── Supabase helpers ────────────────────────────────────────────────────────

def _clean_filename(name: str) -> str:
    name = str(name or "").strip().replace(" ", "_")
    name = re.sub(r"[^a-zA-Z0-9._-]", "", name)
    return name or "file"


def _jr_folder(jr_no: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+', "_", str(jr_no)) if jr_no else "pending"


# ─── Routes ─────────────────────────────────────────────────────────────────

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "nlp_loaded": NLP_LOADED})


@app.route("/api/jr-master", methods=["GET"])
def get_jr_master():
    """Fetch active JR numbers with skill and job details."""
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"error": "Supabase not configured"}), 500
    try:
        resp = requests.get(
            f"{SUPABASE_URL}/rest/v1/jr_master"
            "?select=jr_no,skill_name,job_details,jr_status"
            "&jr_status=eq.active"
            "&order=jr_no.asc",
            headers=_headers(),
            timeout=10,
        )
        if resp.status_code != 200:
            return jsonify({"error": f"Supabase error {resp.status_code}"}), 500
        return jsonify(resp.json()), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/parse-resume", methods=["POST"])
def parse_resume_endpoint():
    """
    Parse a resume file (PDF/DOCX) using the real Python parser.
    Returns: first_name, last_name, email, phone, country, confidence
    """
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    allowed = f.filename.lower().endswith((".pdf", ".docx"))
    if not allowed:
        return jsonify({"error": "Only PDF and DOCX supported"}), 400

    try:
        file_bytes = f.read()
        result = parse_resume_bytes(file_bytes, f.filename)
        return jsonify(result), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e), "first_name": "", "last_name": "",
                        "email": "", "phone": "", "confidence": "low"}), 200


@app.route("/api/upload-resume", methods=["POST"])
def upload_resume():
    """Upload resume file to Supabase Storage."""
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"error": "Supabase not configured"}), 500

    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    jr_no = request.form.get("jr_no", "")

    try:
        file_bytes = f.read()
        file_hash  = hashlib.md5(file_bytes).hexdigest()[:8]
        safe_name  = _clean_filename(f.filename)
        folder     = _jr_folder(jr_no)
        file_path  = f"{folder}/{file_hash}_{safe_name}"

        url = f"{SUPABASE_URL}/storage/v1/object/resumes/{file_path}"
        resp = requests.post(url, headers=_headers(is_binary=True), data=file_bytes, timeout=30)

        if resp.status_code == 409:
            # Already exists — path is still correct
            return jsonify({"path": file_path, "existed": True}), 200
        if resp.status_code not in (200, 201):
            return jsonify({"error": f"Storage upload failed: {resp.text}"}), 400

        return jsonify({"path": file_path, "existed": False}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/check-duplicate", methods=["POST"])
def check_duplicate():
    """Check if a candidate already exists for a given JR."""
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"exists": False}), 200
    try:
        data  = request.get_json()
        jr_no = data.get("jr_no", "")
        email = data.get("email", "")
        phone = data.get("phone", "")

        filters = [f"jr_number=eq.{jr_no}"]
        if email:
            filters.append(f"email=eq.{email}")
        if phone:
            filters.append(f"phone=eq.{phone}")

        query = "&".join(filters) + "&select=id,email,phone&limit=1"
        resp  = requests.get(
            f"{SUPABASE_URL}/rest/v1/candidates_submitted?{query}",
            headers=_headers(),
            timeout=10,
        )
        rows = resp.json() if resp.status_code == 200 else []
        return jsonify({"exists": bool(rows), "id": rows[0].get("id") if rows else None}), 200
    except Exception as e:
        return jsonify({"exists": False, "error": str(e)}), 200


@app.route("/api/submit-candidates", methods=["POST"])
def submit_candidates():
    """Insert candidate records into Supabase and trigger GitHub workflow."""
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"error": "Supabase not configured"}), 500

    try:
        data            = request.get_json()
        candidates      = data.get("candidates", [])
        jr_no           = data.get("jr_no", "").strip()
        recruiter_email = data.get("recruiter_email", "").strip()
        skill           = data.get("skill", "").strip()
        client_recruiter       = data.get("client_recruiter", "")
        client_recruiter_email = data.get("client_recruiter_email", "")

        if not jr_no or not recruiter_email:
            return jsonify({"error": "Missing JR number or recruiter email"}), 400
        if not candidates:
            return jsonify({"error": "No candidates to submit"}), 400
        for cand in candidates:
            missing = []
            if not (cand.get("first_name") or "").strip():
                missing.append("first name")
            if not (cand.get("last_name") or "").strip():
                missing.append("last name")
            if not (cand.get("email") or "").strip():
                missing.append("email")
            if not (cand.get("phone") or "").strip():
                missing.append("phone")
            if not (cand.get("resume_path") or "").strip():
                missing.append("resume path")
            if missing:
                label = (cand.get("file_name") or "candidate").strip()
                return jsonify({"error": f"{label} is missing required fields: {', '.join(missing)}"}), 400

        today_text      = datetime.now(timezone.utc).strftime("%d-%b-%Y")
        recruiter_name  = _name_from_email(recruiter_email)
        inserted_ids    = []
        summary         = []

        for cand in candidates:
            email = (cand.get("email") or "").strip()
            phone = (cand.get("phone") or "").strip()
            first = (cand.get("first_name") or "").strip()
            last  = (cand.get("last_name") or "").strip()

            payload = {
                "jr_number":               jr_no,
                "date_text":               today_text,
                "skill":                   skill,
                "file_name":               cand.get("file_name", ""),
                "first_name":              first,
                "last_name":               last,
                "email":                   email,
                "phone":                   phone,
                "resume_path":             cand.get("resume_path", ""),
                "upload_to_sap":           "Pending",
                "recruiter":               recruiter_name,
                "recruiter_email":         recruiter_email,
                "client_recruiter":        client_recruiter,
                "client_recruiter_email":  client_recruiter_email,
                "actual_status":           "Not Called",
                "call_iteration":          "First Call",
                "created_by":              recruiter_email,
                "modified_by":             recruiter_email,
            }

            resp = requests.post(
                f"{SUPABASE_URL}/rest/v1/candidates_submitted",
                headers=_headers(),
                json=payload,
                timeout=30,
            )

            if resp.status_code in (200, 201):
                rec = resp.json()
                rid = (rec[0].get("id") if isinstance(rec, list) and rec else rec.get("id", ""))
                inserted_ids.append(str(rid))
                summary.append({"candidate": f"{first} {last}".strip() or cand.get("file_name"), "status": "Queued ✓"})
            elif resp.status_code == 409 or "23505" in resp.text:
                summary.append({"candidate": f"{first} {last}".strip() or cand.get("file_name"), "status": "Duplicate"})
            else:
                summary.append({"candidate": f"{first} {last}".strip() or cand.get("file_name"), "status": f"Error {resp.status_code}"})

        # Trigger GitHub Actions
        gh_ok = False
        if inserted_ids and GH_REPO and GH_TOKEN:
            try:
                gh_resp = requests.post(
                    f"https://api.github.com/repos/{GH_REPO}/dispatches",
                    headers={
                        "Authorization":        f"Bearer {GH_TOKEN}",
                        "Accept":               "application/vnd.github+json",
                        "X-GitHub-Api-Version": "2022-11-28",
                        "Content-Type":         "application/json",
                    },
                    json={
                        "event_type": GH_EVENT,
                        "client_payload": {
                            "record_ids":      inserted_ids,
                            "recruiter_email": recruiter_email,
                            "submitted_at":    datetime.now(timezone.utc).isoformat(),
                        },
                    },
                    timeout=15,
                )
                gh_ok = gh_resp.status_code in (200, 201, 204)
            except Exception:
                pass

        return jsonify({
            "success":      True,
            "count":        len(inserted_ids),
            "github_fired": gh_ok,
            "summary":      summary,
            "message":      f"✓ {len(inserted_ids)} candidate(s) queued for SAP upload. "
                            f"{'Email notification will be sent.' if gh_ok else 'Check scheduled workflow for processing.'}",
        }), 200

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


def _name_from_email(email: str) -> str:
    if not email or "@" not in email:
        return ""
    name_part = email.split("@")[0]
    name = name_part.replace(".", " ").replace("-", " ").replace("_", " ")
    return " ".join(w.capitalize() for w in name.split())


if __name__ == "__main__":
    port  = int(os.getenv("PORT", 5000))
    debug = os.getenv("FLASK_ENV") == "development"
    app.run(host="0.0.0.0", port=port, debug=debug)
