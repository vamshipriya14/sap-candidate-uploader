import re
import io
import tempfile
import pdfplumber
import docx
import spacy
import phonenumbers

try:
    import pytesseract
    import pypdfium2 as pdfium
except Exception:
    pytesseract = None
    pdfium = None

nlp = spacy.load("en_core_web_sm")

# =========================
# 🔶 TEXT EXTRACTION
# =========================
def extract_text(file):
    text = ""

    if file.name.endswith(".pdf"):
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for p in pdf.pages:
                page_text = p.extract_text() or p.extract_text(layout=True) or ""
                text += page_text + "\n"

    elif file.name.endswith(".docx"):
        file.seek(0)
        doc = docx.Document(file)
        for p in doc.paragraphs:
            text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

    # Strip private-use Unicode characters (e.g. Wingdings/symbol bullets like \uf097)
    # These appear as line separators in some PDFs and cause lines to merge incorrectly.
    text = re.sub(r'[\ue000-\uf8ff]', '\n', text)

    return text


def _should_try_ocr(text):
    stripped = re.sub(r"\s+", " ", str(text or "")).strip()
    return len(stripped) < 50


def extract_text_via_ocr(file):
    if pytesseract is None or pdfium is None:
        return ""
    if not str(getattr(file, "name", "")).lower().endswith(".pdf"):
        return ""

    try:
        file.seek(0)
        pdf_bytes = file.read()
        if not pdf_bytes:
            return ""

        pdf = pdfium.PdfDocument(pdf_bytes)
        parts = []
        max_pages = min(len(pdf), 5)
        for index in range(max_pages):
            page = pdf[index]
            bitmap = page.render(scale=2.0)
            image = bitmap.to_pil()
            text = pytesseract.image_to_string(image) or ""
            if text.strip():
                parts.append(text)
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_name_from_filename(file_name):
    base = re.sub(r"\.[^.]+$", "", str(file_name or "").strip())
    base = re.sub(r"[_\-]+", " ", base)
    base = re.sub(r"\s+", " ", base).strip()
    words = [w for w in re.split(r"\s+", base) if w]
    words = [w for w in words if w.isalpha() and w.lower() not in INVALID_WORDS]
    if not words:
        return "", ""
    if len(words) == 1:
        return words[0].capitalize(), ""
    pretty = [w.capitalize() for w in words[:4]]
    return " ".join(pretty[:-1]), pretty[-1]


# =========================
# 🔶 EMAIL
# =========================
def extract_email(text):
    # 🔥 Fix broken emails split across lines or spaces
    text = re.sub(r'(\w+)\.\s*\n\s*(\w+@)', r'\1.\2', text)

    # Normalize spaces around @ and .
    text = re.sub(r'(\S+)\s*@\s*(\S+)', r'\1@\2', text)
    text = re.sub(r'(\S+)\s*\.\s*(\S+)', r'\1.\2', text)

    # Fix OCR misreads: comma instead of dot inside email-like patterns
    # e.g. "sureshchandra,ide@gmail.com" -> "sureshchandra.ide@gmail.com"
    text = re.sub(r'(\w),(\w+@)', r'\1.\2', text)          # comma before @
    text = re.sub(r'(@[A-Za-z0-9\-]+),([A-Za-z]{2,})', r'\1.\2', text)  # comma in domain

    emails = re.findall(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', text)

    return emails[0] if emails else ""


# =========================
# 🔶 PHONE
# =========================
def extract_phone(text):
    # Try India first, then fall back to no region hint
    for region in ("IN", None):
        for match in phonenumbers.PhoneNumberMatcher(text, region):
            return phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.E164)
    return ""


# =========================
# 🔶 COUNTRY
# =========================
# Map dial codes → (SAP country code, SAP country name)
COUNTRY_MAP = {
    "+91":  ("+91",  "India"),
    "+1":   ("+1",   "United States"),
    "+44":  ("+44",  "United Kingdom"),
    "+61":  ("+61",  "Australia"),
    "+971": ("+971", "United Arab Emirates"),
    "+65":  ("+65",  "Singapore"),
    "+60":  ("+60",  "Malaysia"),
    "+49":  ("+49",  "Germany"),
    "+33":  ("+33",  "France"),
    "+81":  ("+81",  "Japan"),
}

def extract_country(phone):
    for code, value in COUNTRY_MAP.items():
        if phone.startswith(code):
            return value
    return "+91", "India"   # default


# =========================
# 🔶 NAME EXTRACTION
# =========================
INVALID_WORDS = {
    "resume", "cv", "professional", "summary", "profile", "skills",
    "experience", "business", "operations", "management", "core",
    "competencies", "personal", "details", "id", "objective",
    "education", "contact", "address", "declaration", "references",
    "linkedin", "github", "portfolio",
    # Common resume section headers and filler words that are not names
    "trained", "secondary", "module", "technical", "current", "past",
    "present", "from", "to", "date", "with", "at", "in", "for", "of",
    "and", "the", "by", "as", "on", "an", "or", "is", "was", "are",
    "has", "have", "had", "been", "be", "do", "does", "did",
}

def extract_name(text, email=None):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return "", "", "low"

    def clean_line(line):
        line = re.sub(r'\S+@\S+', '', line)               # remove emails
        line = re.sub(r'\+?\d[\d\s\-]{8,}', '', line)     # remove phone numbers
        line = re.sub(r'(?i)\b(email|mail|contact|phone|mobile)\b[:\-]*', '', line)
        # Split CamelCase single tokens (e.g. "GoureshMathapathi" → "Gouresh Mathapathi")
        line = re.sub(r'([a-z])([A-Z])', r'\1 \2', line)
        line = re.sub(r'[^\w\s]', ' ', line)
        return re.sub(r'\s+', ' ', line).strip()

    def is_valid(words):
        if not (2 <= len(words) <= 4):
            return False
        return all(
            not any(c.isdigit() for c in w) and w.lower() not in INVALID_WORDS
            for w in words
        )

    def split_name(words):
        words = [w.capitalize() for w in words]
        # Single initial at start: treat rest as first name, initial as last
        if len(words[0]) == 1:
            return " ".join(words[1:]), words[0]
        if len(words) == 2:
            return words[0], words[1]
        # 3+ words: everything except last word = first name
        return " ".join(words[:-1]), words[-1]

    # Run spaCy once on first 1000 chars
    spacy_doc = nlp(text[:1000])
    spacy_names = {ent.text.lower() for ent in spacy_doc.ents if ent.label_ == "PERSON"}

    def spacy_match(name):
        name_lower = name.lower()
        return any(name_lower in s or s in name_lower for s in spacy_names)

    # Build candidate list
    candidates = []
    for idx, line in enumerate(lines[:20]):
        words = clean_line(line).split()
        if is_valid(words):
            candidates.append((words, idx))

    # Boost candidates near the email line
    for i, line in enumerate(lines):
        if "@" in line:
            for target_line in ([lines[i - 1]] if i > 0 else []) + [line]:
                words = clean_line(target_line).split()
                if is_valid(words):
                    candidates.append((words, i))

    # Score candidates
    email_username = re.sub(r'[^a-z]', '', email.split("@")[0].lower()) if email else ""
    best_words, best_score = None, -1

    for words, position in candidates:
        name = " ".join(words)
        score = 2                               # base: valid format

        if position < 5:   score += 2          # near top of resume
        elif position < 10: score += 1

        if 2 <= len(words) <= 3: score += 2    # ideal name length

        if email_username and any(w.lower() in email_username for w in words):
            score += 3                          # name appears in email

        if spacy_match(name):
            score += 3                          # spaCy recognises as PERSON

        if score > best_score:
            best_score, best_words = score, words

    if best_words:
        first, last = split_name(best_words)
        confidence = "high" if best_score >= 7 else "medium" if best_score >= 4 else "low"
        return first, last, confidence

    return "", "", "low"


# =========================
# 🔶 DOCX → PDF FALLBACK
# =========================
def convert_docx_to_pdf(file) -> str:
    """Convert an uploaded .docx file to a temp PDF. Returns PDF path."""
    import pythoncom
    from docx2pdf import convert

    pythoncom.CoInitialize()

    file.seek(0)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file.read())
        docx_path = tmp.name

    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path


# =========================
# 🔶 MAIN PARSER
# =========================
def parse_resume(file) -> dict:
    text = extract_text(file)
    if _should_try_ocr(text):
        ocr_text = extract_text_via_ocr(file)
        if len(ocr_text.strip()) > len(text.strip()):
            text = ocr_text
    # Fix broken words across newlines (common in PDFs)
    text = re.sub(r'(\S)\n(\S)', r'\1 \2', text)
    email = extract_email(text)
    first, last, confidence = extract_name(text, email)
    phone = extract_phone(text)
    code, country = extract_country(phone)

    # Fallback: convert .docx → PDF and retry if extraction failed
    if (not first or not email) and file.name.endswith(".docx"):
        pdf_path = convert_docx_to_pdf(file)
        with pdfplumber.open(pdf_path) as pdf:
            text = "".join(p.extract_text() or "" for p in pdf.pages)
            # Fix broken words across newlines (common in PDFs)
            text = re.sub(r'(\S)\n(\S)', r'\1 \2', text)

        email = extract_email(text)
        first, last, confidence = extract_name(text, email)
        phone = extract_phone(text)
        code, country = extract_country(phone)

    if not first and not last:
        first, last = extract_name_from_filename(getattr(file, "name", ""))
        if first or last:
            confidence = "low"

    return {
        "first_name":  first,
        "last_name":   last,
        "email":       email,
        "phone":       phone,
        "country_code": code,
        "country":     country,
        "confidence":  confidence,
    }